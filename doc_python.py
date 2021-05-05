#Read docx file
from docx import Document
read_dcox_file  = Document("sample.docx")
content = read_dcox_file.paragraphs
for i in content:
    print(i.text)
    
#=========================================================

#Create docx file
import docx
from docx import *

file = Document()
file.add_heading("Training Schedule",level=1)
file.add_heading("Subject",level=3)
file.add_paragraph("Maths",style="List Number")
file.add_paragraph("Physics",style="List Number")
file.add_paragraph("Chemistry",style="List Number")
file.add_paragraph("History",style="List Number")
file.add_paragraph("Geography",style="List Number")
file.add_paragraph("Bengali",style="List Number")
file.add_paragraph("English",style="List Number")
file.add_heading("Tools", level=1)
file.add_paragraph("WHITE Board",style="List Bullet")
file.add_paragraph("Marker",style="List Bullet")
file.add_paragraph("Skype",style="List Bullet")
file.save('d:/ss.docx')
print("docx file successfully created...!!")
